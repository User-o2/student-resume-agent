"""简历智能体的核心工具函数与 LangChain Tool 封装。"""

from __future__ import annotations

import json
import re
from io import BytesIO
from datetime import datetime
from pathlib import Path
from typing import Any, Mapping

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from jinja2 import Environment, FileSystemLoader, StrictUndefined
from langchain_core.language_models.chat_models import BaseChatModel
from langchain_core.tools import tool

from app.config import DEFAULT_TEMPLATE_PATH, OUTPUTS_DIR
from app.prompts import POLISH_EXPERIENCE_PROMPT
from app.schema import Award, Experience, ResumeState, Skills


STAGE_LABELS = {
    "personal_info": "个人信息",
    "education": "教育背景",
    "projects": "项目经历",
    "awards": "竞赛获奖",
    "self_evaluation": "自我评价",
    "ready": "可生成简历",
}


def parse_json_object(text: str) -> dict[str, Any]:
    """从文本中解析 JSON 对象。

    Args:
        text: 待解析文本。

    Returns:
        JSON 字典；解析失败时返回空字典。
    """

    if not text:
        return {}

    cleaned = text.strip()
    if cleaned.startswith("```"):
        cleaned = re.sub(r"^```(?:json)?", "", cleaned, flags=re.IGNORECASE).strip()
        cleaned = re.sub(r"```$", "", cleaned).strip()

    try:
        parsed = json.loads(cleaned)
        return parsed if isinstance(parsed, dict) else {}
    except json.JSONDecodeError:
        start = cleaned.find("{")
        end = cleaned.rfind("}")
        if start == -1 or end == -1 or end <= start:
            return {}
        try:
            parsed = json.loads(cleaned[start : end + 1])
            return parsed if isinstance(parsed, dict) else {}
        except json.JSONDecodeError:
            return {}


def coerce_resume_state(state: ResumeState | Mapping[str, Any] | str | None) -> ResumeState:
    """将不同输入形式转换为 ResumeState。

    Args:
        state: 简历状态对象、字典、JSON 字符串或 None。

    Returns:
        标准化后的简历状态对象。
    """

    if isinstance(state, ResumeState):
        return state.model_copy(deep=True)
    if isinstance(state, str):
        data = parse_json_object(state)
        return ResumeState.model_validate(data or {})
    if isinstance(state, Mapping):
        return ResumeState.model_validate(dict(state))
    return ResumeState()


def _is_blank(value: Any) -> bool:
    """判断字段值是否为空。

    Args:
        value: 任意字段值。

    Returns:
        是否为空。
    """

    return value is None or value == "" or value == [] or value == {}


def _unique_extend(base: list[Any], patch: list[Any]) -> list[Any]:
    """合并列表并保持顺序去重。

    Args:
        base: 原列表。
        patch: 新列表。

    Returns:
        合并后的列表。
    """

    result = list(base)
    for item in patch:
        if _is_blank(item):
            continue
        if item not in result:
            result.append(item)
    return result


def _split_scalar_list_text(text: str, key: str) -> list[str]:
    """将模型误返回的列表字符串拆成列表。

    Args:
        text: 原始字符串。
        key: 字段名，用于选择拆分策略。

    Returns:
        字符串列表。
    """

    cleaned = text.strip().strip("。；;")
    if not cleaned:
        return []

    if key in {"courses", "programming_languages", "tools", "professional_skills", "languages", "technologies"}:
        parts = re.split(r"[、,，;；/|\n]+", cleaned)
    else:
        bullet_text = re.sub(r"(?:^|\n)\s*[-*]\s*", "\n", cleaned)
        parts = re.split(r"\n+|[；;]+", bullet_text)
    return [part.strip(" -：:。.") for part in parts if part.strip(" -：:。.")]


def _coerce_list_value(value: Any, key: str) -> list[Any]:
    """将字段值规整为列表。

    Args:
        value: 原始字段值。
        key: 字段名。

    Returns:
        列表形式字段值。
    """

    if _is_blank(value):
        return []
    if isinstance(value, list):
        return value
    if isinstance(value, tuple | set):
        return list(value)
    if isinstance(value, str):
        return _split_scalar_list_text(value, key)
    return [value]


def _normalize_resume_payload(payload: Any, key: str = "") -> Any:
    """规整 LLM 或工具输入的 ResumeState 字段形状。

    Args:
        payload: 原始补丁或状态数据。
        key: 当前字段名。

    Returns:
        可被 `ResumeState` 校验的规整数据。
    """

    list_fields = {
        "courses",
        "technologies",
        "responsibilities",
        "results",
        "polished_bullets",
        "programming_languages",
        "tools",
        "professional_skills",
        "languages",
        "highlights",
    }
    record_list_fields = {"projects", "internships", "awards"}

    if key in record_list_fields and isinstance(payload, Mapping):
        payload = [payload]
    if key in list_fields:
        return [_normalize_resume_payload(item) for item in _coerce_list_value(payload, key)]
    if isinstance(payload, Mapping):
        return {item_key: _normalize_resume_payload(item_value, item_key) for item_key, item_value in payload.items()}
    if isinstance(payload, list):
        return [_normalize_resume_payload(item) for item in payload]
    return payload


def _normalize_major_text(value: Any) -> str:
    """规整专业名称，避免模板重复追加“专业”。

    Args:
        value: 原始专业字段值。

    Returns:
        去掉末尾“专业”的专业名称。
    """

    text = str(value or "").strip()
    return text.removesuffix("专业").strip()


def _merge_dict(base: dict[str, Any], patch: Mapping[str, Any]) -> dict[str, Any]:
    """递归合并字典。

    Args:
        base: 原始字典。
        patch: 更新字典。

    Returns:
        合并后的字典。
    """

    for key, value in patch.items():
        if _is_blank(value):
            continue

        if isinstance(base.get(key), dict) and isinstance(value, Mapping):
            base[key] = _merge_dict(dict(base[key]), value)
        elif isinstance(base.get(key), list) and isinstance(value, list):
            base[key] = _unique_extend(list(base[key]), value)
        else:
            base[key] = value
    return base


def _record_key(record: Mapping[str, Any], fallback_key: str) -> str:
    """提取经历或奖项的合并键。

    Args:
        record: 记录字典。
        fallback_key: 默认键名。

    Returns:
        可用于匹配的键值。
    """

    return str(record.get(fallback_key) or record.get("title") or record.get("name") or "").strip()


def _record_has_content(record: Mapping[str, Any]) -> bool:
    """判断记录是否包含有效内容。

    Args:
        record: 记录字典。

    Returns:
        是否包含非空字段。
    """

    ignored_keys = {"created_at", "updated_at"}
    for key, value in record.items():
        if key in ignored_keys:
            continue
        if isinstance(value, Mapping):
            if _record_has_content(value):
                return True
        elif isinstance(value, list):
            if any(not _is_blank(item) for item in value):
                return True
        elif not _is_blank(value):
            return True
    return False


def _merge_records(
    base: list[dict[str, Any]],
    patch: list[Mapping[str, Any]],
    fallback_key: str,
) -> list[dict[str, Any]]:
    """按标题或名称合并列表记录。

    Args:
        base: 原始记录列表。
        patch: 新记录列表。
        fallback_key: 用于匹配的默认键名。

    Returns:
        合并后的记录列表。
    """

    result = [dict(item) for item in base if isinstance(item, Mapping) and _record_has_content(item)]
    for item in patch:
        if not isinstance(item, Mapping) or not _record_has_content(item):
            continue
        item_key = _record_key(item, fallback_key)
        match_index = None
        for index, existed in enumerate(result):
            if item_key and _record_key(existed, fallback_key) == item_key:
                match_index = index
                break
        if match_index is None and not item_key and fallback_key == "title" and result:
            # 用户常以“补充：技术栈/职责/成果”的形式继续完善上一段项目，
            # 这类信息通常没有标题，应合并到最近一段项目而不是新建空标题项目。
            match_index = len(result) - 1
        if match_index is None:
            result.append(dict(item))
        else:
            result[match_index] = _merge_dict(result[match_index], item)
    return result


def collect_resume_info(
    state: ResumeState | Mapping[str, Any] | str | None,
    updates: Mapping[str, Any] | str,
) -> ResumeState:
    """整理并更新用户已提供的简历字段。

    Args:
        state: 当前简历状态。
        updates: 本轮用户输入抽取出的字段更新。

    Returns:
        更新后的简历状态。
    """

    resume_state = coerce_resume_state(state)
    patch = parse_json_object(updates) if isinstance(updates, str) else dict(updates)
    patch = _normalize_resume_payload(patch)
    if not patch:
        return resume_state

    data = resume_state.model_dump()
    for key in ("projects", "internships"):
        if isinstance(patch.get(key), list):
            data[key] = _merge_records(data.get(key, []), patch[key], "title")
            patch = {item_key: item_value for item_key, item_value in patch.items() if item_key != key}
    if isinstance(patch.get("awards"), list):
        data["awards"] = _merge_records(data.get("awards", []), patch["awards"], "name")
        patch = {item_key: item_value for item_key, item_value in patch.items() if item_key != "awards"}

    data = _merge_dict(data, patch)
    data = _normalize_resume_payload(data)
    data["projects"] = _merge_records([], data.get("projects", []), "title")
    data["internships"] = _merge_records([], data.get("internships", []), "title")
    data["awards"] = _merge_records([], data.get("awards", []), "name")

    if data["basic_info"].get("university") and not data["education"].get("school"):
        data["education"]["school"] = data["basic_info"]["university"]
    if data["basic_info"].get("major"):
        data["basic_info"]["major"] = _normalize_major_text(data["basic_info"]["major"])
    if data["education"].get("major"):
        data["education"]["major"] = _normalize_major_text(data["education"]["major"])
    if data["basic_info"].get("major") and not data["education"].get("major"):
        data["education"]["major"] = data["basic_info"]["major"]
    if data.get("skills", {}).get("languages") and not data["education"].get("english_level"):
        data["education"]["english_level"] = " | ".join(data["skills"]["languages"])

    updated_state = ResumeState.model_validate(data)
    updated_state.touch()
    return updated_state


def _has_required_tech_stack(state: ResumeState) -> bool:
    """判断是否已有可填入模板的技术栈。

    Args:
        state: 当前简历状态。

    Returns:
        是否存在至少一种技术或专业技能。
    """

    project_technologies = [
        technology
        for project in _meaningful_experiences(state.projects)
        for technology in project.technologies
    ]
    return bool(
        state.skills.programming_languages
        or state.skills.tools
        or state.skills.professional_skills
        or project_technologies
    )


def _experience_has_template_content(experience: Experience) -> bool:
    """判断经历是否能支撑模板中的项目条目。

    Args:
        experience: 项目经历。

    Returns:
        是否包含标题和至少一条可展示描述。
    """

    return bool(
        experience.title
        and (
            experience.polished_bullets
            or experience.responsibilities
            or experience.results
            or experience.raw_description
        )
    )


def _award_has_template_content(award: Award) -> bool:
    """判断奖项是否能支撑模板中的竞赛获奖条目。

    Args:
        award: 奖项记录。

    Returns:
        是否包含奖项名称和至少一条可展示描述。
    """

    return bool(award.name and (award.highlights or award.description or award.date or award.level))


def _experience_quality_questions(experience: Experience, category: str) -> list[str]:
    """生成经历完整度追问。

    Args:
        experience: 项目或实习经历。
        category: 经历类别。

    Returns:
        需要追问的问题列表。
    """

    questions: list[str] = []
    title = experience.title or f"这段{category}"
    if not experience.technologies:
        questions.append(f"{title}还缺少技术方法或工具，请补充使用的框架、模型、数据库或软件。")
    if not experience.responsibilities:
        questions.append(f"{title}还缺少个人职责，请补充你具体负责的模块或工作内容。")
    if not experience.results:
        questions.append(f"{title}还缺少项目成果，请补充效果、指标、排名、上线情况或可验证结果。")
    return questions


def _experience_score(experience: Experience) -> int:
    """计算经历信息完整度分数。

    Args:
        experience: 项目或实习经历。

    Returns:
        完整度分数。
    """

    score = 0
    score += 2 if experience.title else 0
    score += 3 if experience.technologies else 0
    score += 3 if experience.responsibilities else 0
    score += 3 if experience.results else 0
    score += 1 if experience.raw_description else 0
    return score


def _meaningful_experiences(experiences: list[Experience]) -> list[Experience]:
    """过滤空经历并按完整度降序返回。

    Args:
        experiences: 原始经历列表。

    Returns:
        有效经历列表。
    """

    meaningful = [experience for experience in experiences if _experience_score(experience) > 0]
    return sorted(meaningful, key=_experience_score, reverse=True)


def _has_valid_phone(text: str) -> bool:
    """判断电话字段是否满足基础格式。

    Args:
        text: 电话字段文本。

    Returns:
        是否为可接受的电话格式。
    """

    return bool(re.fullmatch(r"(?:1[3-9]\d{9}|\d{3,4}[-\s]?\d{7,8})", text.strip()))


def _has_valid_email(text: str) -> bool:
    """判断邮箱字段是否满足基础格式。

    Args:
        text: 邮箱字段文本。

    Returns:
        是否为可接受的邮箱格式。
    """

    return bool(re.fullmatch(r"[\w.+-]+@[\w-]+(?:\.[\w-]+)+", text.strip()))


def validate_resume_state(state: ResumeState | Mapping[str, Any] | str | None) -> dict[str, Any]:
    """执行模板生成前的底线校验。

    Args:
        state: 当前简历状态。

    Returns:
        包含缺失字段、格式错误、质量建议和是否可生成的报告。
    """

    resume_state = coerce_resume_state(state)
    missing_fields: list[str] = []
    validation_errors: list[str] = []
    quality_questions: list[str] = []

    if not resume_state.job_intention.target_position:
        missing_fields.append("个人信息：目标岗位")
    if not resume_state.job_intention.target_industry:
        missing_fields.append("个人信息：目标行业")
    if not resume_state.job_intention.expected_city:
        missing_fields.append("个人信息：期望城市")
    if not resume_state.basic_info.name:
        missing_fields.append("个人信息：姓名")
    if not resume_state.basic_info.phone:
        missing_fields.append("个人信息：电话")
    elif not _has_valid_phone(resume_state.basic_info.phone):
        validation_errors.append("个人信息：电话格式不正确")
    if not resume_state.basic_info.email:
        missing_fields.append("个人信息：邮箱")
    elif not _has_valid_email(resume_state.basic_info.email):
        validation_errors.append("个人信息：邮箱格式不正确")
    if not resume_state.basic_info.native_place:
        missing_fields.append("个人信息：籍贯")
    if not (resume_state.education.school or resume_state.basic_info.university):
        missing_fields.append("教育背景：学校")
    if not resume_state.education.college:
        missing_fields.append("教育背景：学院")
    if not (resume_state.education.major or resume_state.basic_info.major):
        missing_fields.append("教育背景：专业")
    if not resume_state.education.gpa_or_rank:
        missing_fields.append("教育背景：专业排名")
    if not resume_state.education.english_level:
        missing_fields.append("教育背景：英语水平")
    if not resume_state.education.courses:
        missing_fields.append("教育背景：核心课程")
    if not _has_required_tech_stack(resume_state):
        missing_fields.append("教育背景：技术栈")

    meaningful_projects = _meaningful_experiences(resume_state.projects)
    if meaningful_projects:
        if not _experience_has_template_content(meaningful_projects[0]):
            missing_fields.append("项目经历：项目标题和项目要点")
        else:
            quality_questions.extend(_experience_quality_questions(meaningful_projects[0], "项目经历"))
    else:
        missing_fields.append("项目经历：至少 1 段项目经历")

    meaningful_awards = [award for award in resume_state.awards if _award_has_template_content(award)]
    if not meaningful_awards:
        missing_fields.append("竞赛获奖：至少 1 项竞赛、奖学金或证书")

    if not resume_state.self_evaluation:
        missing_fields.append("自我评价：2-3 条个人优势")

    blocking_fields = missing_fields + validation_errors
    return {
        "is_ready": not blocking_fields,
        "can_generate": not blocking_fields,
        "required_missing_fields": missing_fields,
        "missing_fields": missing_fields,
        "validation_errors": validation_errors,
        "blocking_fields": blocking_fields,
        "quality_questions": quality_questions,
        "optional_suggestions": [],
        "current_stage": resume_state.current_stage,
    }


def check_missing_fields(state: ResumeState | Mapping[str, Any] | str | None) -> dict[str, Any]:
    """兼容旧接口的简历底线校验函数。

    Args:
        state: 当前简历状态。

    Returns:
        包含缺失字段、质量建议和是否可生成的报告。
    """

    return validate_resume_state(state)


def _clean_markdown_list(items: list[str]) -> list[str]:
    """清洗 Markdown 列表条目。

    Args:
        items: 原始条目列表。

    Returns:
        清洗后的条目列表。
    """

    cleaned: list[str] = []
    for item in items:
        text = str(item).strip().lstrip("-").strip()
        if re.fullmatch(r"(?:\d+[.、]\s*)?(技术栈|个人职责|项目成果|项目结果|成果)", text):
            continue
        if text and text not in cleaned:
            cleaned.append(text)
    return cleaned


def polish_experience(
    raw_text: str,
    target_position: str = "",
    llm: BaseChatModel | None = None,
) -> list[str]:
    """将口语化经历润色为简历要点。

    Args:
        raw_text: 原始经历描述。
        target_position: 目标岗位。
        llm: 可选的聊天模型实例。

    Returns:
        简历化表达的要点列表。
    """

    text = raw_text.strip()
    if not text:
        return []

    if llm is not None:
        prompt = POLISH_EXPERIENCE_PROMPT.format(
            target_position=target_position or "学生求职/实习",
            raw_text=text,
        )
        try:
            response = llm.invoke(prompt)
            content = getattr(response, "content", str(response))
            bullets = _clean_markdown_list(content.splitlines())
            if bullets:
                return bullets[:3]
        except Exception:
            pass

    normalized = re.sub(r"\s+", " ", text).strip("，。；; ")
    normalized = normalized.replace("我主要", "主要").replace("我负责", "负责")
    if not normalized.startswith(("参与", "负责", "主导", "协助", "完成")):
        normalized = f"参与{normalized}"

    bullets = [f"{normalized}，沉淀了与{target_position or '目标岗位'}相关的实践经验。"]
    if not re.search(r"\d|%|提升|准确率|排名|上线|用户|效率", text):
        bullets.append("建议继续补充量化成果，例如准确率、性能提升、用户规模、完成模块数量或竞赛排名。")
    return bullets


def polish_state_experiences(
    state: ResumeState | Mapping[str, Any] | str | None,
    llm: BaseChatModel | None = None,
    force: bool = False,
) -> ResumeState:
    """批量润色状态中的项目与实习经历。

    Args:
        state: 当前简历状态。
        llm: 可选的聊天模型实例。
        force: 是否覆盖已有的 polished_bullets。

    Returns:
        润色后的简历状态。
    """

    resume_state = coerce_resume_state(state)
    target = resume_state.job_intention.target_position

    for group in (resume_state.projects, resume_state.internships):
        for experience in group:
            if experience.polished_bullets and not force:
                continue
            if llm is None and (experience.responsibilities or experience.results):
                experience.polished_bullets = _clean_markdown_list(
                    experience.responsibilities + experience.results
                )
                continue
            raw_parts = [
                experience.raw_description,
                "；".join(experience.responsibilities),
                "；".join(experience.results),
            ]
            raw_text = "；".join(part for part in raw_parts if part)
            experience.polished_bullets = polish_experience(raw_text, target, llm)
    resume_state.touch()
    return resume_state


def _join_or_default(items: list[str], default: str = "待补充", separator: str = "、") -> str:
    """将列表拼接为指定分隔符文本。

    Args:
        items: 字符串列表。
        default: 空列表时的默认文本。
        separator: 拼接使用的分隔符。

    Returns:
        拼接后的文本。
    """

    return separator.join(items) if items else default


def _format_experiences(experiences: list[Experience], empty_text: str, max_items: int | None = None) -> str:
    """格式化项目经历为新版模板 Markdown。

    Args:
        experiences: 经历列表。
        empty_text: 空列表时的展示文本。
        max_items: 最多展示的经历数量。

    Returns:
        Markdown 文本。
    """

    meaningful_experiences = _meaningful_experiences(experiences)
    if not meaningful_experiences:
        return empty_text

    blocks: list[str] = []
    selected_experiences = meaningful_experiences[:max_items] if max_items else meaningful_experiences
    for experience in selected_experiences:
        title = experience.title or "未命名经历"
        lines = [f"**{title}**"]
        bullets = experience.polished_bullets or (experience.responsibilities + experience.results)
        if not bullets and experience.raw_description:
            bullets = polish_experience(experience.raw_description)
        lines.extend(f"- {item}" for item in _clean_markdown_list(bullets))
        blocks.append("\n".join(lines))
    return "\n\n".join(blocks)


def _format_awards(awards: list[Award]) -> str:
    """格式化竞赛获奖为新版模板 Markdown。

    Args:
        awards: 奖项列表。

    Returns:
        Markdown 文本。
    """

    if not awards:
        return "- 待补充"

    blocks: list[str] = []
    for award in awards:
        if not _award_has_template_content(award):
            continue
        lines = [f"**{award.name or '未命名奖项'}**"]
        highlights = list(award.highlights)
        if not highlights:
            detail_parts = [part for part in [award.date, award.level, award.description] if part]
            if detail_parts:
                highlights = [" | ".join(detail_parts)]
        lines.extend(f"- {item}" for item in _clean_markdown_list(highlights))
        blocks.append("\n".join(lines))
    return "\n\n".join(blocks) if blocks else "- 待补充"


def _format_self_evaluation(text: str) -> str:
    """格式化自我评价为 Markdown 列表。

    Args:
        text: 自我评价原始文本。

    Returns:
        Markdown 列表文本。
    """

    if not text.strip():
        return "- 待补充"

    items = _clean_markdown_list(re.split(r"[\n。；;]+", text))
    return "\n".join(f"- {item}" for item in items[:4]) if items else f"- {text.strip()}"


def _format_tech_stack(skills: Skills, projects: list[Experience] | None = None) -> str:
    """格式化模板教育背景中的技术栈。

    Args:
        skills: 技能结构化字段。
        projects: 项目经历列表，用于补充项目技术栈。

    Returns:
        技术栈文本。
    """

    items = (
        list(skills.programming_languages)
        + list(skills.tools)
        + list(skills.professional_skills)
    )
    for project in _meaningful_experiences(projects or []):
        items.extend(project.technologies)
    deduped: list[str] = []
    for item in items:
        if item and item not in deduped and not re.search(r"(专业排名|英语水平|核心课程|主修课程|GPA)", item):
            deduped.append(item)
    return _join_or_default(deduped, separator=", ")


def build_template_context(state: ResumeState | Mapping[str, Any] | str | None) -> dict[str, str]:
    """构建 Jinja2 简历模板上下文。

    Args:
        state: 当前简历状态。

    Returns:
        模板上下文字典。
    """

    resume_state = coerce_resume_state(state)
    basic = resume_state.basic_info
    job = resume_state.job_intention
    education = resume_state.education
    skills = resume_state.skills

    return {
        "name": basic.name or "姓名待补充",
        "target_position": job.target_position or "待补充",
        "target_industry": job.target_industry or "待补充",
        "expected_city": job.expected_city or "待补充",
        "phone": basic.phone or "待补充",
        "email": basic.email or "待补充",
        "native_place": basic.native_place or "待补充",
        "university": basic.university or education.school or "待补充",
        "college": education.college or "待补充",
        "major": basic.major or education.major or "待补充",
        "education_header": (
            f"{education.school or basic.university or '待补充'} "
            f"{education.college or '待补充'} "
            f"{education.major or basic.major or '待补充'}专业  "
        ),
        "rank": education.gpa_or_rank or "待补充",
        "english_level": education.english_level or "待补充",
        "core_courses": _join_or_default(education.courses, separator="，"),
        "tech_stack": _format_tech_stack(skills, resume_state.projects),
        "projects": _format_experiences(resume_state.projects, "- 待补充", max_items=3),
        "awards": _format_awards(resume_state.awards),
        "self_evaluation": _format_self_evaluation(resume_state.self_evaluation),
        "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
    }


def fill_resume_template(
    state: ResumeState | Mapping[str, Any] | str | None,
    template_path: Path | str = DEFAULT_TEMPLATE_PATH,
    output_path: Path | str | None = None,
) -> dict[str, str]:
    """将结构化简历状态填入 Markdown 模板并保存。

    Args:
        state: 当前简历状态。
        template_path: Markdown 模板路径。
        output_path: 可选输出路径。

    Returns:
        包含 Markdown 内容和输出路径的字典。
    """

    template_file = Path(template_path)
    if not template_file.exists():
        raise FileNotFoundError(f"简历模板不存在：{template_file}")

    env = Environment(
        loader=FileSystemLoader(str(template_file.parent)),
        undefined=StrictUndefined,
        autoescape=False,
        trim_blocks=True,
        lstrip_blocks=True,
    )
    template = env.get_template(template_file.name)
    markdown = template.render(**build_template_context(state)).strip() + "\n"

    destination = Path(output_path) if output_path else OUTPUTS_DIR / f"resume_{datetime.now():%Y%m%d_%H%M%S}.md"
    destination.parent.mkdir(parents=True, exist_ok=True)
    destination.write_text(markdown, encoding="utf-8")

    return {"markdown": markdown, "output_path": str(destination)}


def _set_word_run_font(run: Any, font_size: int | None = None, bold: bool | None = None) -> None:
    """统一设置 Word 文本的中文字体和基础样式。

    Args:
        run: python-docx 的文本运行对象。
        font_size: 可选字号，单位为磅。
        bold: 可选加粗状态。

    Returns:
        None。
    """

    run.font.name = "Microsoft YaHei"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "Microsoft YaHei")
    if font_size is not None:
        run.font.size = Pt(font_size)
    if bold is not None:
        run.bold = bold


def _append_markdown_inline(paragraph: Any, text: str) -> None:
    """将 Markdown 行内加粗语法写入 Word 段落。

    Args:
        paragraph: python-docx 的段落对象。
        text: 待写入的 Markdown 文本。

    Returns:
        None。
    """

    for fragment in re.split(r"(\*\*[^*]+\*\*)", text.rstrip()):
        if not fragment:
            continue
        is_bold = fragment.startswith("**") and fragment.endswith("**")
        run = paragraph.add_run(fragment[2:-2] if is_bold else fragment)
        _set_word_run_font(run, bold=is_bold)


def _create_word_document(markdown_text: str) -> Document:
    """将项目生成的简历 Markdown 转换为排版清晰的 Word 文档。

    Args:
        markdown_text: 已完成模板渲染的简历 Markdown 内容。

    Returns:
        可保存的 Word 文档对象。
    """

    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    normal_style = document.styles["Normal"]
    normal_style.font.size = Pt(10.5)
    normal_style.font.name = "Microsoft YaHei"
    normal_style.element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "Microsoft YaHei")

    for raw_line in markdown_text.splitlines():
        line = raw_line.rstrip()
        if not line:
            continue

        if line.startswith("# "):
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(10)
            run = paragraph.add_run(line[2:].strip())
            _set_word_run_font(run, font_size=18, bold=True)
            continue

        if line.startswith("## "):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(8)
            paragraph.paragraph_format.space_after = Pt(4)
            run = paragraph.add_run(line[3:].strip())
            _set_word_run_font(run, font_size=13, bold=True)
            continue

        if line.startswith("### "):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(3)
            run = paragraph.add_run(line[4:].strip())
            _set_word_run_font(run, font_size=11, bold=True)
            continue

        if line.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.paragraph_format.space_after = Pt(1)
            _append_markdown_inline(paragraph, line[2:].strip())
            continue

        title_match = re.fullmatch(r"\*\*(.+?)\*\*", line.strip())
        if title_match:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(4)
            paragraph.paragraph_format.space_after = Pt(1)
            run = paragraph.add_run(title_match.group(1))
            _set_word_run_font(run, font_size=11, bold=True)
            continue

        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        _append_markdown_inline(paragraph, line.strip())

    return document


def export_resume_to_word(
    markdown_text: str,
    output_path: Path | str | None = None,
) -> dict[str, str | bytes]:
    """将已生成的简历 Markdown 导出为 Word 文档。

    Args:
        markdown_text: 已生成的 Markdown 简历内容。
        output_path: 可选的 .docx 输出路径。

    Returns:
        包含 Word 二进制内容和输出路径的字典。

    Raises:
        ValueError: Markdown 内容为空或输出路径不是 .docx 文件时抛出。
    """

    if not markdown_text.strip():
        raise ValueError("Markdown 简历为空，无法导出 Word 文件。")

    destination = Path(output_path) if output_path else OUTPUTS_DIR / f"resume_{datetime.now():%Y%m%d_%H%M%S}.docx"
    if destination.suffix.lower() != ".docx":
        raise ValueError("Word 输出路径必须以 .docx 结尾。")

    document = _create_word_document(markdown_text)
    buffer = BytesIO()
    document.save(buffer)
    document_bytes = buffer.getvalue()

    destination.parent.mkdir(parents=True, exist_ok=True)
    destination.write_bytes(document_bytes)
    return {"docx_bytes": document_bytes, "output_path": str(destination)}


@tool
def collect_resume_info_tool(current_state_json: str, update_json: str) -> str:
    """整理并更新用户已提供的简历字段。

    Args:
        current_state_json: 当前简历状态 JSON。
        update_json: 本轮字段更新 JSON。

    Returns:
        更新后的简历状态 JSON。
    """

    state = collect_resume_info(current_state_json, update_json)
    return state.model_dump_json(ensure_ascii=False)


@tool
def validate_resume_state_tool(current_state_json: str) -> str:
    """执行模板生成前的底线校验。

    Args:
        current_state_json: 当前简历状态 JSON。

    Returns:
        缺失字段、格式错误与质量建议 JSON 报告。
    """

    report = validate_resume_state(current_state_json)
    return json.dumps(report, ensure_ascii=False)


@tool
def check_missing_fields_tool(current_state_json: str) -> str:
    """兼容旧工具名的底线校验工具。

    Args:
        current_state_json: 当前简历状态 JSON。

    Returns:
        缺失字段、格式错误与质量建议 JSON 报告。
    """

    return validate_resume_state_tool.invoke({"current_state_json": current_state_json})


@tool
def polish_experience_tool(raw_text: str, target_position: str = "") -> str:
    """对项目或实习经历进行简历化润色。

    Args:
        raw_text: 原始经历描述。
        target_position: 目标岗位。

    Returns:
        Markdown 要点文本。
    """

    return "\n".join(f"- {item}" for item in polish_experience(raw_text, target_position))


@tool
def fill_resume_template_tool(current_state_json: str) -> str:
    """将结构化简历状态填入模板并保存 Markdown 简历。

    Args:
        current_state_json: 当前简历状态 JSON。

    Returns:
        包含 Markdown 内容和输出路径的 JSON。
    """

    result = fill_resume_template(current_state_json)
    preview = "\n".join(result["markdown"].splitlines()[:12])
    return json.dumps(
        {
            "markdown": "",
            "output_path": result["output_path"],
            "preview": preview,
        },
        ensure_ascii=False,
    )


RESUME_TOOLS = [
    collect_resume_info_tool,
    validate_resume_state_tool,
    fill_resume_template_tool,
]
