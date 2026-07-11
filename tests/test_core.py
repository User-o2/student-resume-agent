"""核心工具函数的单元测试。"""

from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from docx import Document

from app.schema import ResumeState
from app.tools import (
    check_missing_fields,
    collect_resume_info,
    export_resume_to_word,
    fill_resume_template,
    polish_experience,
)


class ResumeToolTestCase(unittest.TestCase):
    """测试简历工具函数的确定性行为。"""

    def test_collect_resume_info_merges_nested_fields(self) -> None:
        """验证嵌套字段和列表记录可以正确合并。"""

        state = ResumeState()
        state = collect_resume_info(
            state,
            {
                "basic_info": {"name": "李明", "university": "浙江工业大学"},
                "projects": [{"title": "校园平台", "technologies": ["Python"]}],
            },
        )
        state = collect_resume_info(
            state,
            {
                "basic_info": {"email": "liming@example.com"},
                "projects": [{"title": "校园平台", "results": ["完成 18 个接口"]}],
            },
        )

        self.assertEqual(state.basic_info.name, "李明")
        self.assertEqual(state.basic_info.email, "liming@example.com")
        self.assertEqual(len(state.projects), 1)
        self.assertEqual(state.projects[0].technologies, ["Python"])
        self.assertEqual(state.projects[0].results, ["完成 18 个接口"])

    def test_check_missing_fields_reports_project_quality_without_blocking_generation(self) -> None:
        """验证项目质量追问不会阻塞已满足模板必填板块的简历生成。"""

        state = collect_resume_info(
            ResumeState(),
            {
                "basic_info": {
                    "name": "李明",
                    "university": "浙江工业大学",
                    "major": "计算机科学与技术",
                    "phone": "13800000001",
                    "email": "liming@example.com",
                    "native_place": "山东省济南市",
                },
                "job_intention": {
                    "target_position": "Python 后端实习",
                    "target_industry": "互联网",
                    "expected_city": "杭州",
                },
                "education": {
                    "school": "浙江工业大学",
                    "college": "计算机学院",
                    "major": "计算机科学与技术",
                    "courses": ["数据结构"],
                    "gpa_or_rank": "专业前 15%",
                    "english_level": "CET-6 510",
                },
                "projects": [{"title": "图像识别项目", "raw_description": "我做过一个图像识别项目"}],
                "skills": {"programming_languages": ["Python"]},
                "awards": [{"name": "校级程序设计竞赛二等奖", "description": "团队排名前 10%"}],
                "self_evaluation": "学习能力强，关注后端开发方向。",
            },
        )

        report = check_missing_fields(state)
        self.assertTrue(report["is_ready"])
        self.assertTrue(report["quality_questions"])

    def test_check_missing_fields_blocks_only_required_fields(self) -> None:
        """验证新版模板所有必填板块缺失时会阻塞生成。"""

        state = collect_resume_info(
            ResumeState(),
            {
                "basic_info": {
                    "name": "李明",
                    "university": "浙江工业大学",
                    "major": "计算机科学与技术",
                },
                "job_intention": {"target_position": "Python 后端实习"},
            },
        )

        report = check_missing_fields(state)

        self.assertFalse(report["is_ready"])
        self.assertIn("个人信息：目标行业", report["missing_fields"])
        self.assertIn("个人信息：电话", report["missing_fields"])
        self.assertIn("个人信息：邮箱", report["missing_fields"])
        self.assertIn("个人信息：籍贯", report["missing_fields"])
        self.assertIn("教育背景：学院", report["missing_fields"])
        self.assertIn("项目经历：至少 1 段项目经历", report["missing_fields"])
        self.assertIn("竞赛获奖：至少 1 项竞赛、奖学金或证书", report["missing_fields"])
        self.assertIn("自我评价：2-3 条个人优势", report["missing_fields"])

    def test_collect_resume_info_merges_untitled_project_supplements(self) -> None:
        """验证无标题项目补充会合并到上一段项目。"""

        state = collect_resume_info(
            ResumeState(),
            {
                "projects": [
                    {
                        "technologies": ["Flask"],
                        "responsibilities": ["负责智能体路由模块"],
                    }
                ]
            },
        )
        state = collect_resume_info(
            state,
            {
                "projects": [
                    {
                        "technologies": ["Redis", "Celery"],
                        "results": ["接口吞吐量提升约 40%"],
                    }
                ]
            },
        )

        self.assertEqual(len(state.projects), 1)
        self.assertEqual(state.projects[0].technologies, ["Flask", "Redis", "Celery"])
        self.assertEqual(state.projects[0].results, ["接口吞吐量提升约 40%"])

    def test_collect_resume_info_normalizes_llm_string_lists(self) -> None:
        """验证 LLM 将列表字段误返回为字符串时可以自动规整。"""

        state = collect_resume_info(
            ResumeState(),
            {
                "education": {"courses": "机器学习、深度学习、数据结构与算法"},
                "projects": {
                    "title": "基于轻量级CNN的手写数字识别系统",
                    "technologies": "TensorFlow, Keras, Flask",
                    "responsibilities": "负责数据清洗与增强，使用TensorFlow/Keras搭建MobileNetV2变体，另封装为Flask接口。",
                    "results": "测试集准确率99.1%；响应时间低于200ms。",
                },
                "skills": {
                    "programming_languages": "Python, C++",
                    "tools": "PyTorch, Linux, Git",
                },
            },
        )

        self.assertEqual(state.education.courses, ["机器学习", "深度学习", "数据结构与算法"])
        self.assertEqual(state.projects[0].technologies, ["TensorFlow", "Keras", "Flask"])
        self.assertEqual(
            state.projects[0].responsibilities,
            ["负责数据清洗与增强，使用TensorFlow/Keras搭建MobileNetV2变体，另封装为Flask接口"],
        )
        self.assertEqual(state.projects[0].results, ["测试集准确率99.1%", "响应时间低于200ms"])
        self.assertEqual(state.skills.programming_languages, ["Python", "C++"])
        self.assertEqual(state.skills.tools, ["PyTorch", "Linux", "Git"])

    def test_collect_resume_info_normalizes_major_suffix(self) -> None:
        """验证专业字段末尾的“专业”会被规整，避免模板重复。"""

        state = collect_resume_info(
            ResumeState(),
            {
                "basic_info": {"major": "人工智能专业"},
                "education": {"major": "人工智能专业"},
            },
        )

        self.assertEqual(state.education.major, "人工智能")
        self.assertNotIn("major", state.basic_info.model_dump())

    def test_legacy_basic_education_fields_migrate_and_can_correct_state(self) -> None:
        """验证旧版学校、专业字段可读取，并能作为补丁修正规范字段。"""

        state = ResumeState.model_validate(
            {
                "basic_info": {
                    "name": "李明",
                    "university": "旧学校",
                    "major": "旧专业",
                },
                "education": {
                    "school": "规范学校",
                    "major": "规范专业",
                },
            }
        )

        self.assertEqual(state.education.school, "规范学校")
        self.assertEqual(state.education.major, "规范专业")
        self.assertNotIn("university", state.basic_info.model_dump())
        self.assertNotIn("major", state.basic_info.model_dump())

        corrected = collect_resume_info(
            state,
            {
                "basic_info": {
                    "university": "浙江工业大学",
                    "major": "计算机科学与技术专业",
                }
            },
        )

        self.assertEqual(corrected.education.school, "浙江工业大学")
        self.assertEqual(corrected.education.major, "计算机科学与技术")

    def test_check_missing_fields_uses_best_meaningful_project(self) -> None:
        """验证空项目不会导致完整项目被误判为缺技术。"""

        state = collect_resume_info(
            ResumeState(),
            {
                "basic_info": {
                    "name": "小明",
                    "university": "天津科技大学",
                    "major": "人工智能",
                    "phone": "15208898345",
                    "email": "27485938@qq.com",
                    "native_place": "天津市",
                },
                "job_intention": {
                    "target_position": "Python 后端实习",
                    "target_industry": "互联网",
                    "expected_city": "杭州",
                },
                "education": {
                    "school": "天津科技大学",
                    "college": "人工智能学院",
                    "major": "人工智能",
                    "courses": ["机器学习"],
                    "gpa_or_rank": "3/105",
                    "english_level": "CET-6 500",
                },
                "projects": [
                    {},
                    {
                        "title": "智能体路由系统",
                        "technologies": ["Flask", "Redis", "Celery"],
                        "responsibilities": ["负责智能体路由与对话管理模块"],
                        "results": ["接口吞吐量提升约 40%"],
                    },
                ],
                "skills": {"programming_languages": ["Python"]},
                "awards": [{"name": "校级程序设计竞赛二等奖", "description": "团队排名前 10%"}],
                "self_evaluation": "有较好的开发经验，认真学习新技术。",
            },
        )

        report = check_missing_fields(state)

        self.assertTrue(report["is_ready"])
        self.assertEqual(report["quality_questions"], [])

    def test_fill_resume_template_writes_markdown(self) -> None:
        """验证模板填充可以写出 Markdown 文件。"""

        state = collect_resume_info(
            ResumeState(),
            {
                "basic_info": {
                    "name": "李明",
                    "university": "浙江工业大学",
                    "major": "计算机科学与技术",
                    "phone": "13800000001",
                    "email": "liming@example.com",
                    "native_place": "山东省济南市",
                },
                "job_intention": {
                    "target_position": "Python 后端实习",
                    "target_industry": "互联网",
                    "expected_city": "杭州",
                },
                "education": {
                    "school": "浙江工业大学",
                    "college": "计算机学院",
                    "major": "计算机科学与技术",
                    "courses": ["数据结构"],
                    "gpa_or_rank": "专业前 15%",
                    "english_level": "CET-6 510",
                },
                "projects": [
                    {
                        "title": "校园平台",
                        "technologies": ["Python"],
                        "responsibilities": ["负责后端接口开发"],
                        "results": ["完成 18 个接口"],
                    }
                ],
                "skills": {"programming_languages": ["Python"]},
                "awards": [{"name": "校级程序设计竞赛二等奖", "description": "团队排名前 10%"}],
                "self_evaluation": "具备后端开发实践经验。",
            },
        )

        with tempfile.TemporaryDirectory() as tmp_dir:
            output_path = Path(tmp_dir) / "resume.md"
            result = fill_resume_template(state, output_path=output_path)

        self.assertIn("# 李明", result["markdown"])
        self.assertIn("校园平台", result["markdown"])
        self.assertIn("## 竞赛获奖", result["markdown"])
        self.assertTrue(Path(result["output_path"]).name.endswith("resume.md"))

    def test_fill_resume_template_renders_internship_details_and_word(self) -> None:
        """验证实习经历完整进入 Markdown，并可继续导出 Word。"""

        state = collect_resume_info(
            ResumeState(),
            {
                "basic_info": {"name": "李明"},
                "internships": [
                    {
                        "title": "教务数据管理实践",
                        "organization": "学校信息化办公室",
                        "role": "后端开发实践",
                        "start_date": "2025.03",
                        "end_date": "2025.06",
                        "responsibilities": ["使用 Python 与 Pandas 清洗课程成绩数据"],
                        "results": ["将人工整理时间从半天缩短到约 30 分钟"],
                    }
                ],
            },
        )

        with tempfile.TemporaryDirectory() as tmp_dir:
            markdown_path = Path(tmp_dir) / "resume.md"
            word_path = Path(tmp_dir) / "resume.docx"
            rendered = fill_resume_template(state, output_path=markdown_path)
            export_resume_to_word(rendered["markdown"], output_path=word_path)
            document = Document(word_path)

        self.assertIn("## 实习经历", rendered["markdown"])
        self.assertIn("教务数据管理实践", rendered["markdown"])
        self.assertIn("学校信息化办公室 | 后端开发实践 | 2025.03 - 2025.06", rendered["markdown"])
        self.assertIn("将人工整理时间从半天缩短到约 30 分钟", rendered["markdown"])
        word_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        self.assertIn("实习经历", word_text)
        self.assertIn("教务数据管理实践", word_text)

    def test_fill_resume_template_uses_internship_note_when_empty(self) -> None:
        """验证没有实习记录时在模板中展示用户提供的说明。"""

        state = ResumeState(internship_note="暂无正式实习经历，项目实践作为主要证明材料。")

        with tempfile.TemporaryDirectory() as tmp_dir:
            result = fill_resume_template(state, output_path=Path(tmp_dir) / "resume.md")

        self.assertIn("## 实习经历", result["markdown"])
        self.assertIn("暂无正式实习经历，项目实践作为主要证明材料。", result["markdown"])

    def test_export_resume_to_word_writes_readable_docx(self) -> None:
        """验证已生成的 Markdown 简历可以导出为可读取的 Word 文件。"""

        markdown = """# 李明

- **求职意向**：Python 后端实习 | 互联网 | 杭州
- **电话**：13800000001

## 教育背景
浙江工业大学 计算机学院 计算机科学与技术专业
- **核心课程**：数据结构，计算机网络

## 项目经历
**校园平台**
- 负责后端接口开发，完成用户认证与商品管理模块。
"""

        with tempfile.TemporaryDirectory() as tmp_dir:
            output_path = Path(tmp_dir) / "resume.docx"
            result = export_resume_to_word(markdown, output_path=output_path)

            self.assertTrue(output_path.exists())
            self.assertGreater(len(result["docx_bytes"]), 0)
            document = Document(output_path)
            content = "\n".join(paragraph.text for paragraph in document.paragraphs)

        self.assertEqual(result["output_path"], str(output_path))
        self.assertIn("李明", content)
        self.assertIn("教育背景", content)
        self.assertIn("校园平台", content)
        self.assertIn("负责后端接口开发", content)

    def test_export_resume_to_word_rejects_empty_markdown(self) -> None:
        """验证空简历不会被导出为 Word 文件。"""

        with self.assertRaises(ValueError):
            export_resume_to_word("   ")

    def test_polish_experience_fallback_returns_bullets(self) -> None:
        """验证无 LLM 时经历润色仍能返回简历要点。"""

        bullets = polish_experience("我做过校园网站，主要写后端接口", "Python 后端实习")

        self.assertGreaterEqual(len(bullets), 1)
        self.assertIn("Python 后端实习", bullets[0])


if __name__ == "__main__":
    unittest.main()
