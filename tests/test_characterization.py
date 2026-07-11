"""核心用户流程的特征测试，保护重构前已经稳定的外部行为。"""

from __future__ import annotations

import json
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

from docx import Document

from app.agent import ResumeAgentService
from app.config import EXAMPLES_DIR
from app.schema import ResumeState
from app.tools import export_resume_to_word


def load_ready_state() -> ResumeState:
    """读取仓库中的完整学生案例。

    Args:
        无。

    Returns:
        可直接生成简历的结构化状态。
    """

    data = json.loads((EXAMPLES_DIR / "student_case_1.json").read_text(encoding="utf-8"))
    return ResumeState.model_validate(data)


class CoreFlowCharacterizationTestCase(unittest.TestCase):
    """固定对话、生成、导入、评分和 Word 导出的现有行为。"""

    def test_resume_state_carries_business_facts_between_service_instances(self) -> None:
        """验证业务事实由 ResumeState 传递，不依赖同一个服务实例的历史。"""

        first_service = ResumeAgentService(use_llm=False, use_agent_driver=False)
        first_result = first_service.handle_message(
            "目标岗位：Python 后端实习；目标行业：互联网；期望城市：杭州；姓名：李明；电话：13800000001；邮箱：liming@example.com；籍贯：山东济南",
            ResumeState(),
        )

        second_service = ResumeAgentService(use_llm=False, use_agent_driver=False)
        second_result = second_service.handle_message(
            "学校：浙江工业大学；学院：计算机学院；专业：计算机科学与技术",
            first_result.state,
        )

        self.assertEqual(second_result.state.basic_info.name, "李明")
        self.assertEqual(second_result.state.job_intention.target_position, "Python 后端实习")
        self.assertEqual(second_result.state.education.school, "浙江工业大学")

    def test_generation_requires_explicit_intent_and_returns_rendered_result(self) -> None:
        """验证完整状态不会自动生成，明确请求后才返回渲染结果。"""

        state = load_ready_state()
        service = ResumeAgentService(use_llm=False, use_agent_driver=False)

        waiting_result = service.handle_message("我再确认一下当前信息", state)
        self.assertEqual(waiting_result.resume_markdown, "")
        self.assertEqual(waiting_result.output_path, "")

        with patch(
            "app.agent._fill_template_with_tool",
            return_value={"markdown": "# 李明\n\n## 项目经历\n", "output_path": "/tmp/resume.md"},
        ):
            generated_result = service.handle_message("生成简历", waiting_result.state)

        self.assertEqual(generated_result.output_path, "/tmp/resume.md")
        self.assertTrue(generated_result.resume_markdown.startswith("# 李明"))

    def test_existing_resume_import_preserves_source_facts(self) -> None:
        """验证已有 Markdown 简历导入后保留姓名、岗位和项目事实。"""

        markdown = """# 王欣

- **求职意向**：机器学习实习 | 人工智能 | 上海
- **电话**：13800000002
- **邮箱**：wangxin@example.com
- **籍贯**：江苏南京

## 教育背景
南京理工大学 人工智能学院 人工智能专业
- **专业排名**：专业前20%
- **英语水平**：CET-4 531
- **核心课程**：机器学习，深度学习
- **技术栈**：Python, PyTorch

## 项目经历
**垃圾分类图像识别模型**
- 负责数据清洗和模型训练
- 验证集准确率达到91.3%

## 竞赛获奖
**数学建模竞赛省级三等奖**
- 负责特征构造与模型实现

## 自我评价
- 具备机器学习项目经验
"""
        service = ResumeAgentService(use_llm=False, use_agent_driver=False)

        with tempfile.TemporaryDirectory() as tmp_dir:
            output_path = Path(tmp_dir) / "optimized.md"

            def fake_fill(
                state: ResumeState,
                trace: list[str],
                output_path: Path | str | None = None,
            ) -> dict[str, str]:
                """模拟模板编排，隔离导入解析行为。

                Args:
                    state: 解析并润色后的状态。
                    trace: Agent 执行轨迹。
                    output_path: 测试输出路径。

                Returns:
                    模拟的模板渲染结果。
                """

                trace.append("调用工具：fill_resume_template")
                destination = Path(output_path) if output_path else Path(tmp_dir) / "optimized.md"
                rendered = f"# {state.basic_info.name}\n\n## 项目经历\n{state.projects[0].title}\n"
                destination.write_text(rendered, encoding="utf-8")
                return {"markdown": rendered, "output_path": str(destination)}

            with patch("app.agent._fill_template_with_tool", side_effect=fake_fill):
                result = service.optimize_existing_resume(markdown, output_path=output_path)

            self.assertTrue(output_path.exists())

        self.assertEqual(result.state.basic_info.name, "王欣")
        self.assertEqual(result.state.job_intention.target_position, "机器学习实习")
        self.assertEqual(result.state.projects[0].title, "垃圾分类图像识别模型")

    def test_offline_scoring_keeps_deterministic_weights(self) -> None:
        """验证离线评分的完整度、匹配度、表达分和综合权重保持稳定。"""

        service = ResumeAgentService(use_llm=False, use_agent_driver=False)
        result = service.score_resume(load_ready_state())

        self.assertEqual(result.report.completeness_score, 100)
        self.assertEqual(result.report.match_score, 93)
        self.assertEqual(result.report.expression_score, 72)
        self.assertEqual(result.report.total_score, 89)

    def test_word_export_preserves_headings_bullets_and_inline_bold(self) -> None:
        """验证 Word 导出保留标题、列表和 Markdown 行内加粗文本。"""

        markdown = """# 李明

- **求职意向**：Python 后端实习 | 互联网 | 杭州

## 项目经历
**校园平台**
- 负责后端接口开发
"""
        with tempfile.TemporaryDirectory() as tmp_dir:
            output_path = Path(tmp_dir) / "resume.docx"
            result = export_resume_to_word(markdown, output_path=output_path)
            document = Document(output_path)

        paragraphs = [paragraph.text for paragraph in document.paragraphs]
        self.assertEqual(result["output_path"], str(output_path))
        self.assertIn("李明", paragraphs)
        self.assertIn("项目经历", paragraphs)
        self.assertIn("校园平台", paragraphs)
        self.assertIn("负责后端接口开发", paragraphs)
        self.assertTrue(
            any(
                run.bold and run.text == "求职意向"
                for paragraph in document.paragraphs
                for run in paragraph.runs
            )
        )


if __name__ == "__main__":
    unittest.main()
